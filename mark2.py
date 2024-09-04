import streamlit as st
import sqlite3
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
from docx import Document
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain.docstore.document import Document as LangchainDocument
import re
import time
from functools import lru_cache
import google.generativeai as genai
from PIL import Image

# Cargar variables de entorno y configurar la API de Google
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Diccionario para rastrear el uso de canciones
song_usage = {}

# Función para obtener embeddings (cacheada)
@st.cache_resource
def get_embeddings():
    return GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# Conexión a la base de datos (cacheada)
@lru_cache(maxsize=None)
def get_db_connection():
    conn = sqlite3.connect('songs_database.db')
    conn.execute('''CREATE TABLE IF NOT EXISTS songs
                    (id INTEGER PRIMARY KEY, name TEXT UNIQUE, content TEXT, theme TEXT)''')
    return conn

# Funciones de base de datos
def song_exists(conn, name):
    return conn.execute("SELECT 1 FROM songs WHERE name = ?", (name,)).fetchone() is not None

def save_song(conn, name, content, theme):
    conn.execute("INSERT OR REPLACE INTO songs (name, content, theme) VALUES (?, ?, ?)", (name, content, theme))
    conn.commit()

def get_all_songs(conn):
    return conn.execute("SELECT name, content, theme FROM songs ORDER BY name ASC").fetchall()

def delete_song(conn, name):
    conn.execute("DELETE FROM songs WHERE name = ?", (name,))
    conn.commit()

# Funciones para procesar archivos
@st.cache_data
def get_file_text(file):
    if file.name.lower().endswith('.pdf'):
        return get_pdf_text(file)
    elif file.name.lower().endswith('.docx'):
        return get_docx_text(file)
    return ""

@st.cache_data
def get_pdf_text(pdf_file):
    text = ""
    pdf_reader = PdfReader(pdf_file)
    for i, page in enumerate(pdf_reader.pages):
        text += f"\n--- Página {i+1} del documento {pdf_file.name} ---\n"
        text += page.extract_text()
    return text

@st.cache_data
def get_docx_text(docx_file):
    doc = Document(docx_file)
    return "\n".join(para.text for para in doc.paragraphs)

# Función para procesar imagen y extraer texto
def process_image_to_text(image_file):
    image = Image.open(image_file)
    model = genai.GenerativeModel('gemini-1.5-pro')
    response = model.generate_content([
        "Lee la imagen, es una canción que va a tener sus respectivas notas en cifrado americano, donde 'C' es 'DO' y 'B' es 'SI'. Presta atención a palabras como 'INTRO' y lo que sigue porque eso indica dónde y cómo empieza la canción. Generalmente, el formato en que vas a ver las canciones es en dos columnas, de izquierda a derecha. Quiero la canción escrita tal cual la foto, sin agregados.",
        image
    ])
    return response.text

# Funciones para análisis de canciones y creación de playlists
@lru_cache(maxsize=1)
def get_song_analysis_chain():
    prompt_template = """
    Analiza el contenido de esta canción y determina su tema principal.
    Proporciona un breve resumen del tema y el tono emocional de la canción.

    Contenido de la canción:
    {context}

    Tema y análisis:
    """
    model = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

def analyze_song(content):
    chain = get_song_analysis_chain()
    doc = LangchainDocument(page_content=content)
    return chain.invoke({"input_documents": [doc], "question": ""})["output_text"]

@lru_cache(maxsize=1)
def get_playlist_creation_chain():
    prompt_template = """
    Basándote en los temas de las canciones proporcionadas y el tema solicitado por el usuario, 
    crea una lista de 8 canciones coherente. Selecciona las canciones que mejor se ajusten al tema solicitado,
    priorizando aquellas que se han usado con menos frecuencia recientemente.
    Explica brevemente por qué estas canciones van bien juntas 
    y cómo se relacionan con el tema solicitado.

    IMPORTANTE: 
    - Asegúrate de no repetir canciones en la playlist. Cada canción debe ser única.
    - Intenta incluir canciones que no se hayan usado recientemente en otras playlists.
    - La playlist DEBE contener exactamente 8 canciones diferentes.

    Tema solicitado por el usuario: {user_theme}

    Canciones disponibles, sus temas y su frecuencia de uso reciente (menor número indica uso menos frecuente):
    {context}

    Lista de reproducción (8 canciones):
    """

    model = ChatGoogleGenerativeAI(
        model="gemini-1.5-pro",
        system_instructions="Eres un experto en música cristiana, del sector reformado, no católico. Crees en la soberanía de Dios, entiendes la Biblia como la palabra de Dios y crees en Jesús como el Hijo de Dios que murió en la Cruz y resucitó al tercer día. Entiendes que las canciones no deben divagar en su temática y tienes que poner La Palabra de Dios en lo más alto",
        temperature=0.4
    )
    prompt = PromptTemplate(template=prompt_template, input_variables=["user_theme", "context"])    
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

# Funciones para manejo de uso de canciones
def update_song_usage(song_name):
    song_usage[song_name] = time.time()

def get_song_usage_score(song_name):
    if song_name not in song_usage:
        return 0
    return 1 / (1 + (time.time() - song_usage[song_name]) / 3600)

# Función para crear playlist
def create_playlist(songs_with_themes, user_theme):
    chain = get_playlist_creation_chain()
    
    context = "\n".join([f"Canción: {song[0]}, Tema: {song[2]}, Uso reciente: {get_song_usage_score(song[0]):.2f}" 
                         for song in songs_with_themes])
    
    doc = LangchainDocument(page_content=context)
    response = chain.invoke(
        {"input_documents": [doc], "user_theme": user_theme, "question": ""}
    )
    playlist = response["output_text"]
    
    song_names = re.findall(r'\d+\.\s*([^,\n]+)', playlist)
    if len(song_names) != 8 or len(song_names) != len(set(song_names)):
        return create_playlist(songs_with_themes, user_theme + " (diverso)")
    
    for song in song_names:
        update_song_usage(song.strip())
    
    return playlist

# Funciones para manejo de canciones en la interfaz
def get_song_content(conn, song_name):
    return conn.execute("SELECT content FROM songs WHERE name = ?", (song_name,)).fetchone()[0]

def display_song_in_columns(song_content, key_prefix):
    lines = song_content.split('\n')
    midpoint = len(lines) // 2
    col1, col2 = st.columns(2)
    with col1:
        st.text_area("Primera parte:", value='\n'.join(lines[:midpoint]), height=400, key=f"{key_prefix}_col1")
    with col2:
        st.text_area("Segunda parte:", value='\n'.join(lines[midpoint:]), height=400, key=f"{key_prefix}_col2")

def update_song_title(conn, old_name, new_name):
    conn.execute("UPDATE songs SET name = ? WHERE name = ?", (new_name, old_name))
    conn.commit()

def update_song_content(conn, name, new_content):
    conn.execute("UPDATE songs SET content = ? WHERE name = ?", (new_content, name))
    conn.commit()

# Función para cambiar la tonalidad de una canción
def change_song_key(content: str, target_key: str) -> str:
    prompt_template = """
    Analiza la siguiente canción con sus notas musicales y cambia todas las notas a la tonalidad de {target_key}.
    Mantén el formato original de la canción, incluyendo la disposición de las notas y la letra.
    Asegúrate de que todas las notas estén correctamente transpuestas a la nueva tonalidad. En el caso de que encuentres que una canción está en notas menores, simplemente respeta la "distancia" entre notas que se te pide y adapta la canción, un ejemplo, si está en 'LA menor' osea 'AM', y te pido que modifiques a 'C' osea que simplemente subas una nota, entonces ese 'AM' debería pasar a 'CM' y así el resto de la canción.
    
    Canción original:
    {context}
    
    Canción en la nueva tonalidad de {target_key}:
    """
    
    model = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.2)
    prompt = PromptTemplate(template=prompt_template, input_variables=["target_key", "context"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    
    doc = LangchainDocument(page_content=content)
    response = chain.invoke({"input_documents": [doc], "target_key": target_key})
    
    return response["output_text"]

# Función para guardar texto como PDF
def save_text_as_pdf(text, filename):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    lines = text.split('\n')
    c.setFont("Helvetica", 12)
    y = height - 50
    
    for line in lines:
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line)
        y -= 15
    
    c.save()
    buffer.seek(0)
    
    with open(filename, "wb") as output_file:
        output_file.write(buffer.getvalue())

def display_songs(conn, search_term=""):
    songs = get_all_songs(conn)
    if songs:
        st.subheader("Canciones Disponibles")
        
        # Barra de búsqueda
        search_term = st.text_input("Buscar canción:", value=search_term)
        
        # Filtrar canciones basado en el término de búsqueda
        if search_term:
            search_pattern = re.compile(r'.*'.join(re.escape(term.lower()) for term in search_term.split()), re.IGNORECASE)
            filtered_songs = [song for song in songs if search_pattern.search(song[0])]
        else:
            filtered_songs = songs
        
        if filtered_songs:
            selected_song = st.selectbox("Selecciona una canción para ver sus detalles:", 
                                         options=[song[0] for song in filtered_songs],
                                         key="song_selector")
            
            if selected_song:
                song_content = get_song_content(conn, selected_song)
                if song_content:
                    st.subheader("Contenido de la canción")
                    
                    # Opción para editar el título de la canción
                    new_title = st.text_input("Título de la canción:", value=selected_song)
                    if new_title != selected_song and st.button("Actualizar título"):
                        update_song_title(conn, selected_song, new_title)
                        st.success(f"Título actualizado de '{selected_song}' a '{new_title}'.")
                        st.rerun()
                    
                    # Opción para editar el contenido de la canción
                    edited_content = st.text_area("Editar contenido de la canción:", value=song_content, height=400)
                    if edited_content != song_content and st.button("Guardar cambios en el contenido"):
                        update_song_content(conn, selected_song, edited_content)
                        st.success("Contenido de la canción actualizado.")
                    
                    # Cambiar tonalidad de la canción
                    st.subheader("Cambiar tonalidad de la canción")
                    target_key = st.selectbox("Selecciona la nueva tonalidad:", 
                                              options=["C", "C#", "D", "D#", "E", "F", "F#", "G", "G#", "A", "A#", "B"])
                    if st.button("Cambiar tonalidad"):
                        with st.spinner("Cambiando tonalidad de la canción... Por favor, espera."):
                            new_song_content = change_song_key(edited_content, target_key)
                            st.subheader(f"Canción en nueva tonalidad: {target_key}")
                            display_song_in_columns(new_song_content, "new_song_content")
                    
                    # Opción para eliminar la canción
                    if st.button(f"Eliminar '{selected_song}'"):
                        delete_song(conn, selected_song)
                        st.success(f"La canción '{selected_song}' ha sido eliminada.")
                        st.rerun()
                else:
                    st.warning("No se pudo cargar el contenido de la canción.")
            
            st.info(f"Mostrando {len(filtered_songs)} de {len(songs)} canciones.")
        else:
            st.info("No se encontraron canciones que coincidan con la búsqueda.")
    else:
        st.info("No hay canciones en la base de datos. Por favor, carga nuevas canciones primero.")
# Función principal de la aplicación
def main():
    st.set_page_config(page_title="Analizador de Canciones y Creador de Playlists", layout="wide")

    st.markdown(
        """
        <h1 style="text-align: center; animation: colorChange 5s infinite;">
        Bienvenido soy Mark2 y voy a ayudarte a analizar Canciones y Crear listas para cantar en las reuniones
        </h1>
        <h3 style='text-align: center; color: #888888;'>Carga tus canciones y crea playlists temáticas</h3>
        """,
        unsafe_allow_html=True
    )

    conn = get_db_connection()

    menu = ["Inicio", "Ver Canciones", "Cargar Nuevas Canciones", "Crear Playlist"]
    choice = st.sidebar.selectbox("Menú", menu)

    if choice == "Inicio":
        st.write("Bienvenido a la aplicación de análisis de canciones y creación de playlists. Utiliza el menú de la izquierda para navegar.")

    elif choice == "Ver Canciones":
        display_songs(conn)

    elif choice == "Cargar Nuevas Canciones":
        st.subheader("Cargar nuevas canciones")
        upload_type = st.radio("Selecciona el tipo de carga:", ["Archivo (PDF/Word)", "Imagen"])
        
        uploaded_files = st.file_uploader(
            "Sube archivos" if upload_type == "Archivo (PDF/Word)" else "Sube imágenes de canciones", 
            accept_multiple_files=True, 
            type=['pdf', 'docx'] if upload_type == "Archivo (PDF/Word)" else ['png', 'jpg', 'jpeg']
        )
        
        if st.button("Procesar Archivos"):
            if uploaded_files:
                with st.spinner("Procesando archivos... Por favor, espera."):
                    for file in uploaded_files:
                        if not song_exists(conn, file.name):
                            content = get_file_text(file) if upload_type == "Archivo (PDF/Word)" else process_image_to_text(file)
                            
                            if upload_type != "Archivo (PDF/Word)":
                                pdf_filename = f"{os.path.splitext(file.name)[0]}.pdf"
                                save_text_as_pdf(content, pdf_filename)
                                st.success(f"Canción extraída de la imagen y guardada como {pdf_filename}")
                            
                            theme = analyze_song(content)
                            save_song(conn, file.name, content, theme)
                            st.success(f"Canción '{file.name}' procesada y guardada.")
                        else:
                            st.warning(f"La canción '{file.name}' ya existe en la base de datos.")
                
                st.success("Todas las canciones han sido procesadas. Puedes crear playlists ahora.")
            else:
                st.warning("Por favor, sube al menos un archivo antes de procesar.")

    elif choice == "Crear Playlist":
        st.subheader("Crear Playlist")
        songs = get_all_songs(conn)
        if songs:
            if len(songs) < 8:
                st.warning(f"Actualmente hay solo {len(songs)} canciones en la base de datos. Se necesitan al menos 8 canciones diferentes para crear una playlist. Por favor, carga más canciones.")
            else:
                user_theme = st.text_input("Ingresa el tema para tu playlist:")
                
                if st.button("Generar Playlist"):
                    if user_theme:
                        with st.spinner("Generando playlist de 8 canciones diversas... Esto puede tomar unos momentos."):
                            playlist = create_playlist(songs, user_theme)
                        st.write("Playlist generada:", playlist)
                    else:
                        st.warning("Por favor, ingresa un tema para la playlist antes de generarla.")
        else:
            st.warning("No hay canciones en la base de datos. Por favor, carga nuevas canciones primero.")

if __name__ == "__main__":
    main()
